from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.document import Document as _Document


def iter_block_items(parent):
    """
    Yield paragraphs and tables in the exact order in which
    they appear inside a document or table cell.
    """

    if isinstance(parent, _Document):
        parent_element = parent.element.body

    elif isinstance(parent, _Cell):
        parent_element = parent._tc

    elif type(parent).__name__ in ('_Header', '_Footer'):
        parent_element = parent._element

    else:
        raise ValueError(
            f"Unsupported parent type: {type(parent)}"
        )

    for child in parent_element.iterchildren():

        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)

        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def build_document_map(doc):
    """
    Build the canonical text representation used by BOTH
    detection and redaction.

    Returns:

        full_text
        blocks

    Each block contains the actual Paragraph object belonging
    to the SAME `doc` object.

    Therefore the redactor MUST modify this same document
    instance instead of loading the DOCX again.
    """

    blocks = []
    text_parts = []

    current_position = 0

    # Prevent processing the same physical paragraph twice.
    processed_paragraphs = set()

    def add_paragraph(paragraph):

        nonlocal current_position

        p_el = paragraph._element
        if p_el in processed_paragraphs:
            return
        processed_paragraphs.add(p_el)

        text = paragraph.text or ""

        if not text:
            return

        start = current_position
        end = start + len(text)

        blocks.append(
            {
                "text": text,
                "element": paragraph,
                "type": "paragraph",
                "start": start,
                "end": end,
            }
        )

        text_parts.append(text)

        # One newline separates this block from the next block.
        current_position = end + 1

    def process_table(table):

        for row in table.rows:

            for cell in row.cells:

                # IMPORTANT:
                # Process paragraphs and nested tables in their
                # actual XML order.
                for item in iter_block_items(cell):

                    if isinstance(item, Paragraph):

                        add_paragraph(item)

                    elif isinstance(item, Table):

                        process_table(item)

    # ========================================================
    # PROCESS BODY IN ACTUAL DOCX ORDER
    # ========================================================

    for item in iter_block_items(doc):

        if isinstance(item, Paragraph):

            add_paragraph(item)

        elif isinstance(item, Table):

            process_table(item)

    # ========================================================
    # PROCESS HEADERS AND FOOTERS
    # ========================================================

    processed_hdrftr = set()
    for section in doc.sections:
        for name, container in [("header", section.header), ("footer", section.footer)]:
            if container.is_linked_to_previous:
                continue
            container_el = container._element
            if container_el in processed_hdrftr:
                continue
            processed_hdrftr.add(container_el)

            for item in iter_block_items(container):
                if isinstance(item, Paragraph):
                    add_paragraph(item)
                elif isinstance(item, Table):
                    process_table(item)

    # ========================================================
    # BUILD EXACT TEXT USED BY DETECTORS
    # ========================================================

    full_text = "\n".join(text_parts)

    return full_text, blocks